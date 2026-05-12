from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

DOCX_OUT = Path('/mnt/data/Entregaveis_Sprint1_CrossPlatform_CCR_Motiva.docx')


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(76, 29, 149)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(75, 85, 99)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Arial'
styles['Normal'].font.size = Pt(11)

section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

add_title(
    doc,
    'Sprint 1 - Exploração, Requisitos e Protótipo',
    'Challenge FIAP + CCR Motiva | Disciplina: Cross Platform (JavaScript)'
)

doc.add_paragraph('Documento consolidado com tudo o que precisa ser entregue na Sprint 1: definição da solução, persona, requisitos, stack, arquitetura, estrutura do repositório e handoff do protótipo.')

add_heading(doc, '1. Projeto proposto', 1)
doc.add_paragraph('Nome do projeto: Motiva VerdeCheck Mobile')
doc.add_paragraph('Problema escolhido: apoiar a equipe de conservação na decisão operacional sobre cortar ou não cortar a grama em trechos de rodovia.')
doc.add_paragraph('Recorte adotado: inspeção em campo com registro do trecho, captura de evidência visual e retorno de recomendação clara de intervenção.')

add_heading(doc, '2. Contexto resumido do challenge', 1)
doc.add_paragraph('A proposta foi estruturada considerando o contexto operacional de conservação de rodovias, em que as equipes executam inspeções rotineiras e precisam preservar canteiro central, faixa de domínio e condições de segurança viária. O desafio pede uso de dados e tecnologia para orientar intervenções com maior precisão.')
add_bullets(doc, [
    'Frentes de conservação envolvem operações rotineiras e de emergência.',
    'Há exigências regulatórias ligadas à faixa de domínio, canteiro central, visibilidade e segurança.',
    'A decisão de roçada depende do contexto da área e da condição observada no trecho.',
    'O MVP mobile foca em padronizar a decisão e registrar evidências.'
])

add_heading(doc, '3. Persona principal', 1)
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
rows = [
    ('Nome', 'Lucas Almeida'),
    ('Cargo', 'Operador de Conservação de Faixa de Domínio'),
    ('Idade', '32 anos'),
    ('Objetivo', 'Registrar o trecho e receber uma recomendação objetiva: Cortar ou Não cortar.'),
    ('Dores', 'Subjetividade da inspeção, retrabalho, falta de evidência rápida e histórico disperso.')
]
for i, (k, v) in enumerate(rows):
    table.cell(i, 0).text = k
    table.cell(i, 1).text = v
    set_cell_shading(table.cell(i, 0), 'EDE9FE')

add_heading(doc, '4. Proposta da solução mobile', 1)
add_bullets(doc, [
    'Cadastro rápido do trecho: rodovia, km, sentido, tipo de área e observações.',
    'Captura ou seleção de foto do local.',
    'Envio dos dados para análise.',
    'Retorno visual com decisão: Cortar ou Não cortar.',
    'Armazenamento do histórico para consulta posterior por operador e supervisor.'
])

add_heading(doc, '5. Requisitos funcionais (RF)', 1)
rf_items = [
    'RF01 - permitir autenticação do usuário.',
    'RF02 - iniciar nova análise.',
    'RF03 - registrar rodovia, km, sentido, tipo de área e observações.',
    'RF04 - capturar ou anexar foto do trecho.',
    'RF05 - enviar dados e imagem para análise.',
    'RF06 - exibir recomendação Cortar ou Não cortar.',
    'RF07 - exibir justificativa resumida.',
    'RF08 - salvar histórico local.',
    'RF09 - listar inspeções anteriores.',
    'RF10 - visualizar detalhe da inspeção.'
]
add_bullets(doc, rf_items)

add_heading(doc, '6. Requisitos não funcionais (RNF)', 1)
rnf_items = [
    'RNF01 - interface simples e rápida para uso em campo.',
    'RNF02 - funcionamento em Android e iOS.',
    'RNF03 - navegação fluida e de baixa complexidade.',
    'RNF04 - persistência local mínima para protótipo.',
    'RNF05 - arquitetura preparada para integração futura com API e IA.',
    'RNF06 - auditabilidade: cada análise deve manter imagem, data e contexto.',
    'RNF07 - código organizado e documentação clara.'
]
add_bullets(doc, rnf_items)

add_heading(doc, '7. Stack tecnológica e justificativa', 1)
stack_table = doc.add_table(rows=1, cols=3)
stack_table.style = 'Table Grid'
headers = ['Camada', 'Tecnologia', 'Justificativa']
for idx, h in enumerate(headers):
    stack_table.cell(0, idx).text = h
    set_cell_shading(stack_table.cell(0, idx), 'DDD6FE')
entries = [
    ('Frontend mobile', 'React Native + Expo', 'Acelera a prototipação cross-platform em JavaScript.'),
    ('Navegação', 'React Navigation', 'Organiza fluxo entre Home, Nova análise, Resultado e Histórico.'),
    ('Armazenamento', 'AsyncStorage', 'Suficiente para persistência local do protótipo.'),
    ('Recursos nativos', 'expo-image-picker / expo-camera / expo-location', 'Preparam o app para captura de imagem e geolocalização.'),
    ('IA futura', 'API Python', 'Permite integrar posteriormente o modelo de visão computacional.'),
]
for layer, tech, why in entries:
    row = stack_table.add_row().cells
    row[0].text = layer
    row[1].text = tech
    row[2].text = why

add_heading(doc, '8. Protótipo navegável - telas principais', 1)
proto_table = doc.add_table(rows=1, cols=3)
proto_table.style = 'Table Grid'
for idx, h in enumerate(['Tela', 'Objetivo', 'Ação principal']):
    proto_table.cell(0, idx).text = h
    set_cell_shading(proto_table.cell(0, idx), 'DDD6FE')
proto_entries = [
    ('Login', 'Acesso ao app', 'Entrar'),
    ('Home', 'Resumo e atalhos', 'Iniciar nova análise'),
    ('Nova análise', 'Registrar contexto do trecho', 'Continuar'),
    ('Captura de imagem', 'Registrar evidência visual', 'Tirar foto'),
    ('Processamento', 'Comunicar análise em andamento', 'Aguardar resultado'),
    ('Resultado', 'Mostrar decisão de corte', 'Salvar inspeção'),
    ('Histórico', 'Consultar registros anteriores', 'Abrir detalhe'),
    ('Detalhe da inspeção', 'Ver evidência e justificativa', 'Compartilhar futuramente'),
]
for a, b, c in proto_entries:
    row = proto_table.add_row().cells
    row[0].text = a
    row[1].text = b
    row[2].text = c

add_heading(doc, '9. Estrutura do repositório', 1)
repo_block = doc.add_paragraph()
repo_block.style = 'Intense Quote'
repo_block.add_run(
    'motiva-verdecheck-mobile/\n'
    '├── docs/\n'
    '│   ├── documento-requisitos.md\n'
    '│   ├── arquitetura-mobile.md\n'
    '│   ├── prototipo-figma-handoff.md\n'
    '│   ├── prompt-figma.txt\n'
    '│   └── sprint1-checklist.md\n'
    '├── src/\n'
    '│   ├── components/\n'
    '│   ├── contexts/\n'
    '│   ├── navigation/\n'
    '│   ├── screens/\n'
    '│   └── utils/\n'
    '├── App.js\n'
    '├── app.json\n'
    '├── package.json\n'
    '└── README.md'
)

add_heading(doc, '10. O que já está pronto para entregar', 1)
add_bullets(doc, [
    'README completo do projeto.',
    'Documento de requisitos em Markdown.',
    'Arquitetura mobile resumida.',
    'Handoff completo para reprodução do protótipo no Figma.',
    'Prompt pronto para IA do Figma.',
    'Protótipo base navegável em React Native + Expo.',
    'Checklist final da sprint.'
])

add_heading(doc, '11. O que o grupo deve personalizar antes da entrega final', 1)
add_bullets(doc, [
    'Substituir os nomes dos integrantes no README.',
    'Publicar o repositório no GitHub e inserir o link final.',
    'Criar o protótipo final no Figma com base no handoff e inserir o link.',
    'Tirar prints do protótipo, se o professor pedir evidência visual adicional.'
])

add_heading(doc, '12. Conclusão', 1)
doc.add_paragraph('A Sprint 1 foi resolvida com foco no que a disciplina pede: entendimento do problema, definição da persona, documentação dos requisitos, justificativa da stack e material completo para prototipação. A solução proposta mantém coerência com o challenge e já deixa a base pronta para evoluir para integração com visão computacional nas próximas sprints.')

doc.save(DOCX_OUT)
print(DOCX_OUT)
